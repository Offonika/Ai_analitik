from __future__ import annotations

import hashlib
import re
from collections.abc import Iterator
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.shared import Cm, Pt, RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph

NAVY = "003153"
BLUE = "2F75B5"
PALE_BLUE = "EDF4FB"
ORANGE = "B87922"
GRAY = "595959"
WHITE = "FFFFFF"
BLACK = "262626"


def markdown_sha256(markdown: str) -> str:
    return hashlib.sha256(markdown.encode("utf-8")).hexdigest()


def render_markdown_docx(
    markdown: str,
    output_path: Path,
    *,
    logo_path: Path | None = None,
    branded: bool = False,
    landscape: bool = True,
    brand_name: str = "Шумейко и Партнеры",
    footer_text: str = "Шумейко и Партнеры · AI-аналитик отчетов",
    cover_title: str = "Аналитический отчет по юнит-экономике WB",
    cover_subtitle: str = "",
    source_sha256: str | None = None,
) -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    _setup_document(
        document,
        logo_path=None if branded else logo_path,
        landscape=landscape,
        brand_name=brand_name,
        footer_text=footer_text,
    )
    if source_sha256:
        document.core_properties.comments = f"source_sha256={source_sha256}"
    if branded:
        _add_brand_cover(
            document,
            logo_path=logo_path,
            title=cover_title,
            subtitle=cover_subtitle,
        )

    lines = _body_lines(markdown)
    index = 0
    skipped_cover_title = False
    in_code_fence = False
    while index < len(lines):
        raw_line = lines[index]
        line = raw_line.strip()
        if line.startswith("```"):
            in_code_fence = not in_code_fence
            index += 1
            continue
        if not line:
            index += 1
            continue
        if not in_code_fence and line.startswith("|"):
            table_rows, index = _parse_markdown_table(lines, index)
            _add_docx_table(document, table_rows)
            continue
        if not in_code_fence and line.startswith("# "):
            if branded and not skipped_cover_title:
                skipped_cover_title = True
            else:
                _add_paragraph(
                    document,
                    line[2:],
                    "Title",
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                )
        elif not in_code_fence and line.startswith("## "):
            paragraph = _add_paragraph(document, line[3:], "Heading 1")
            _set_paragraph_border(paragraph)
        elif not in_code_fence and line.startswith("### "):
            _add_paragraph(document, line[4:], "Heading 2")
        elif not in_code_fence and re.match(r"^\d+\.\s+", line):
            _add_paragraph(document, re.sub(r"^\d+\.\s+", "", line), "List Number")
        elif not in_code_fence and re.match(r"^[-*]\s+", line):
            _add_paragraph(document, re.sub(r"^[-*]\s+", "", line), "List Bullet")
        else:
            _add_paragraph(document, line)
        index += 1

    document.save(output_path)
    return output_path


def normalized_markdown_tokens(markdown: str) -> list[str]:
    tokens: list[str] = []
    lines = _body_lines(markdown)
    index = 0
    in_code_fence = False
    while index < len(lines):
        line = lines[index].strip()
        if line.startswith("```"):
            in_code_fence = not in_code_fence
            index += 1
            continue
        if not line:
            index += 1
            continue
        if not in_code_fence and line.startswith("|"):
            rows, index = _parse_markdown_table(lines, index)
            for row in rows:
                tokens.extend(_normalize_token(cell) for cell in row if cell.strip())
            continue
        if not in_code_fence:
            is_heading = bool(re.match(r"^#{1,6}\s+", line))
            line = re.sub(r"^#{1,6}\s+", "", line)
            if not is_heading:
                line = re.sub(r"^\d+\.\s+", "", line)
                line = re.sub(r"^[-*]\s+", "", line)
        normalized = _normalize_token(line)
        if normalized:
            tokens.append(normalized)
        index += 1
    return tokens


def normalized_docx_tokens(path: Path) -> list[str]:
    document = Document(path)
    tokens: list[str] = []
    for block in _iter_document_blocks(document):
        if isinstance(block, Paragraph):
            normalized = _normalize_token(block.text)
            if normalized:
                tokens.append(normalized)
            continue
        for row in block.rows:
            for cell in row.cells:
                normalized = _normalize_token(cell.text)
                if normalized:
                    tokens.append(normalized)
    return tokens


def docx_source_sha256(path: Path) -> str:
    comments = Document(path).core_properties.comments or ""
    prefix = "source_sha256="
    return comments[len(prefix) :] if comments.startswith(prefix) else ""


def _body_lines(markdown: str) -> list[str]:
    lines = markdown.splitlines()
    if not lines or lines[0].strip() != "---":
        return lines
    for index in range(1, len(lines)):
        if lines[index].strip() == "---":
            return lines[index + 1 :]
    return lines


def _normalize_token(value: str) -> str:
    return " ".join(str(value).replace("\xa0", " ").split())


def _iter_document_blocks(document: DocumentObject) -> Iterator[Paragraph | Table]:
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def _setup_document(
    document: DocumentObject,
    *,
    logo_path: Path | None,
    landscape: bool,
    brand_name: str,
    footer_text: str,
) -> None:
    section = document.sections[0]
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Cm(29.7)
        section.page_height = Cm(21.0)
        section.left_margin = Cm(1.2)
        section.right_margin = Cm(1.2)
    else:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.1)

    normal = document.styles["Normal"]
    normal.font.name = "Geologica"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Geologica")
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    for name, size, color in (
        ("Title", 22, NAVY),
        ("Heading 1", 14, NAVY),
        ("Heading 2", 12, ORANGE),
    ):
        style = document.styles[name]
        style.font.name = "Geologica"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Geologica")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(footer_text)
    _set_run_font(run, size=8, color=GRAY)
    _set_paragraph_border(footer, color=BLUE)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if logo_path is not None and logo_path.exists():
        header.add_run().add_picture(str(logo_path), width=Cm(3.6))
    else:
        run = header.add_run(brand_name)
        _set_run_font(run, size=10, bold=True, color=NAVY)
    _set_paragraph_border(header, color=BLUE)


def _add_brand_cover(
    document: DocumentObject,
    *,
    logo_path: Path | None,
    title: str,
    subtitle: str,
) -> None:
    if logo_path is not None and logo_path.exists():
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(str(logo_path), width=Cm(4.6))
    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_paragraph.add_run(title)
    _set_run_font(run, size=22, bold=True, color=NAVY)
    if subtitle:
        subtitle_paragraph = document.add_paragraph()
        subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle_paragraph.add_run(subtitle)
        _set_run_font(run, size=11, color=GRAY)
    _add_brand_badges(document)
    document.add_paragraph()


def _add_brand_badges(document: DocumentObject) -> None:
    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, label in enumerate(
        ("Юнит-экономика WB", "Возвраты и убыточность", "Упущенные продажи")
    ):
        cell = table.cell(0, index)
        _set_cell_shading(cell, PALE_BLUE)
        _set_cell_border(cell, BLUE)
        _set_cell_margins(cell, top=120, bottom=120)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(label)
        _set_run_font(run, size=9, bold=True, color=NAVY)


def _add_paragraph(
    document: DocumentObject,
    value: str,
    style: str | None = None,
    *,
    align: WD_ALIGN_PARAGRAPH | None = None,
) -> Paragraph:
    paragraph = document.add_paragraph(style=style)
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(str(value))
    _set_run_font(run)
    return paragraph


def _parse_markdown_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        cells = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
        if not all(re.fullmatch(r"[:\-\s]+", cell) for cell in cells):
            rows.append(cells)
        index += 1
    return rows, index


def _add_docx_table(document: DocumentObject, rows: list[list[str]]) -> None:
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=column_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        for column_index in range(column_count):
            value = row[column_index] if column_index < len(row) else ""
            cell = table.cell(row_index, column_index)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _set_cell_border(cell)
            _set_cell_margins(cell)
            if row_index == 0:
                _set_cell_shading(cell, NAVY)
            elif row_index % 2 == 0:
                _set_cell_shading(cell, PALE_BLUE)
            paragraph = cell.paragraphs[0]
            if column_index > 0 and row_index > 0:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = paragraph.add_run(value)
            _set_run_font(
                run,
                size=7 if column_count >= 8 else 8,
                bold=row_index == 0,
                color=WHITE if row_index == 0 else BLACK,
            )
    document.add_paragraph()


def _set_run_font(
    run: Any,
    *,
    size: int | None = None,
    bold: bool | None = None,
    color: str | None = None,
) -> None:
    run.font.name = "Geologica"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Geologica")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def _set_cell_shading(cell: Any, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_border(cell: Any, color: str = "D9E2F3") -> None:
    properties = cell._tc.get_or_add_tcPr()
    borders = properties.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def _set_cell_margins(
    cell: Any,
    top: int = 80,
    start: int = 90,
    bottom: int = 80,
    end: int = 90,
) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for key, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = margins.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_paragraph_border(paragraph: Paragraph, color: str = BLUE) -> None:
    properties = paragraph._p.get_or_add_pPr()
    borders = properties.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        properties.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "10")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
